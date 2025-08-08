"""
Document Content Extractor Service

This module provides content extraction from various document types.
"""

import logging
import io
import json
from typing import Optional
import base64
from app.models.document import DocumentType

logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Extract text content from various document types"""
    
    async def extract_content(
        self,
        file_content: bytes,
        document_type: DocumentType,
        mime_type: Optional[str] = None
    ) -> Optional[str]:
        """
        Extract text content from a document
        
        Args:
            file_content: The raw file content
            document_type: The document type
            mime_type: Optional MIME type
            
        Returns:
            Extracted text content or None if extraction fails
        """
        try:
            if document_type == DocumentType.TEXT:
                return self._extract_text(file_content)
            elif document_type == DocumentType.MARKDOWN:
                return self._extract_markdown(file_content)
            elif document_type == DocumentType.JSON:
                return self._extract_json(file_content)
            elif document_type == DocumentType.HTML:
                return self._extract_html(file_content)
            elif document_type == DocumentType.PDF:
                return await self._extract_pdf(file_content)
            elif document_type == DocumentType.CSV:
                return self._extract_csv(file_content)
            elif document_type == DocumentType.DOCX:
                return await self._extract_docx(file_content)
            elif document_type == DocumentType.XLSX:
                return await self._extract_xlsx(file_content)
            else:
                # Try to extract as text for unknown types
                return self._extract_text(file_content)
                
        except Exception as e:
            logger.error(f"Error extracting content from {document_type}: {e}")
            return None
    
    def _extract_text(self, content: bytes) -> str:
        """Extract plain text content"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            # If all fail, decode with errors ignored
            return content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return ""
    
    def _extract_markdown(self, content: bytes) -> str:
        """Extract markdown content (same as text)"""
        return self._extract_text(content)
    
    def _extract_json(self, content: bytes) -> str:
        """Extract and format JSON content"""
        try:
            json_data = json.loads(content.decode('utf-8'))
            # Convert to formatted string
            return json.dumps(json_data, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"Error extracting JSON: {e}")
            # Fall back to text extraction
            return self._extract_text(content)
    
    def _extract_html(self, content: bytes) -> str:
        """Extract text from HTML"""
        try:
            # Try to import BeautifulSoup
            from bs4 import BeautifulSoup
            
            html = content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(html, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text
            text = soup.get_text()
            
            # Break into lines and remove leading/trailing space
            lines = (line.strip() for line in text.splitlines())
            # Break multi-headlines into a line each
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            # Drop blank lines
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            return text
            
        except ImportError:
            logger.warning("BeautifulSoup not installed, extracting HTML as text")
            return self._extract_text(content)
        except Exception as e:
            logger.error(f"Error extracting HTML: {e}")
            return self._extract_text(content)
    
    async def _extract_pdf(self, content: bytes) -> str:
        """Extract text from PDF"""
        try:
            # Try to import PyPDF2
            import PyPDF2
            
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except ImportError:
            logger.warning("PyPDF2 not installed, cannot extract PDF content")
            return ""
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            return ""
    
    def _extract_csv(self, content: bytes) -> str:
        """Extract CSV content as formatted text"""
        try:
            import csv
            import io
            
            text_content = content.decode('utf-8', errors='ignore')
            csv_reader = csv.reader(io.StringIO(text_content))
            
            rows = []
            for row in csv_reader:
                rows.append(' | '.join(row))
            
            return '\n'.join(rows)
            
        except Exception as e:
            logger.error(f"Error extracting CSV: {e}")
            return self._extract_text(content)
    
    async def _extract_docx(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            # Try to import python-docx
            from docx import Document as DocxDocument
            
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)
            
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(' | '.join(row_text))
            
            return '\n'.join(text)
            
        except ImportError:
            logger.warning("python-docx not installed, cannot extract DOCX content")
            return ""
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ""
    
    async def _extract_xlsx(self, content: bytes) -> str:
        """Extract text from XLSX"""
        try:
            # Try to import openpyxl
            import openpyxl
            
            xlsx_file = io.BytesIO(content)
            workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
            
            text = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text.append(f"=== Sheet: {sheet_name} ===")
                
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty cells
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(row_values):
                        text.append(' | '.join(row_values))
            
            return '\n'.join(text)
            
        except ImportError:
            logger.warning("openpyxl not installed, cannot extract XLSX content")
            return ""
        except Exception as e:
            logger.error(f"Error extracting XLSX: {e}")
            return ""
    
    def get_document_type_from_mime(self, mime_type: str) -> DocumentType:
        """Determine document type from MIME type"""
        mime_map = {
            'text/plain': DocumentType.TEXT,
            'text/markdown': DocumentType.MARKDOWN,
            'text/html': DocumentType.HTML,
            'application/json': DocumentType.JSON,
            'text/csv': DocumentType.CSV,
            'application/pdf': DocumentType.PDF,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentType.DOCX,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': DocumentType.XLSX,
        }
        
        # Check for image types
        if mime_type and mime_type.startswith('image/'):
            return DocumentType.IMAGE
        
        return mime_map.get(mime_type, DocumentType.OTHER)